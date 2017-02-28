import column_name_formatter as cf
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

import ntpath

import os

from PyQt4.QtCore import QObject, pyqtSignal


class DocTable:
    def __init__(self, table_ref):
        self._table = table_ref
        self.cell_style = None

        header_cells = self._table.rows[0].cells[1:]
        self.headers = [h.text for h in header_cells]
        self._standard_headers = [cf.fmt(h) for h in self.headers]
        self.years = []

        self._column_header_map = {cf.fmt(h): idx + 1 for idx, h in enumerate(self.headers)}

        self._year_map = {}
        self._generate_mappings()

    def _generate_mappings(self):
        rows = self._table.rows
        if rows and len(rows) > 1:
            rows = rows[1:]
            for idx, row in enumerate(rows):
                year = row.cells[0].text
                self.years.append(year)
                self._year_map[year] = row.cells

    def set_value(self, year, standard_column_header, value):
        """
        Set a table cells value
        """
        cells = self._year_map[year]
        column = self._column_header_map[standard_column_header]
        cell = cells[column]

        cell.text = value
        p = cell.paragraphs[0]
        p.paragraph_format.alignment = 1
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        if self.cell_style:
            p.style = self.cell_style

    def debug_print(self):
        print(self.headers)


class DocTemplate(QObject):
    """
    Maintains a list of column headings and years to table refs
    """

    started = pyqtSignal()
    finished = pyqtSignal()

    def __init__(self):
        QObject.__init__(self)
        self._doc = None
        self._styles = None
        self.tables = []
        self._table_map = {}
        self.name = ""
        self.loaded = False
        self.all_headers = []
        self.path = ""

    def load(self, path):
        self.started.emit()

        if not path:
            self.loaded = False
        else:
            self.path = path
            self._doc = Document(self.path)
            self.name = ntpath.basename(self.path)

            self._styles = self._doc.styles
            self._add_style("Title", is_bold=True, size=18)
            self._add_style("CellStyle", is_bold=False, size=16)

            self._table_map = {}
            self.tables = []

            self._init_tables()

            self.loaded = True

        self.finished.emit()

    def has_column(self, column):
        return cf.fmt(column) in [cf.fmt(h) for h in self.all_headers]

    def get_all_headers(self):
        return self.all_headers

    def _add_style(self, name, is_bold, size):
        if self._styles and name not in self._styles:
            charstyle = self._styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            font = charstyle.font
            font.bold = is_bold
            font.size = Pt(size)
            font.name = "Arial"

    def _init_tables(self):
        tables = self._doc.tables
        for table in tables:
            doc_table = DocTable(table)
            self.tables.append(doc_table)

            doc_table.cell_style = self._doc.styles["CellStyle"]

            for h in doc_table.headers:
                self._table_map[cf.fmt(h)] = doc_table
                self.all_headers.append(h)

    def set_title(self, constituent):
        if self._doc:
            paragraphs = self._doc.paragraphs
            p = paragraphs[0]

            new_p = p.insert_paragraph_before(constituent)
            new_p.alignment = 1
            new_p.style = self._doc.styles["Title"]

    def write_data(self, year, column_header, value):
        standard_name = cf.fmt(column_header)
        if standard_name in self._table_map:
            table = self._table_map[standard_name]
            table.set_value(year, standard_name, value)

    def save(self, path):
        self._doc.save(path)
